import sys
import json
from docx import Document
import io
import os
from langdetect import detect_langs, DetectorFactory

# Đảm bảo kết quả ổn định khi detect
DetectorFactory.seed = 0

# Đảm bảo output UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Mapping ISO-639-1 -> locale phổ biến
LANG_MAP = {
    "vi": "vi-VN",
    "en": "en-US",
    "fr": "fr-FR",
    "de": "de-DE",
    "es": "es-ES",
    "it": "it-IT",
    "pt": "pt-PT",
    "ru": "ru-RU",
    "ja": "ja-JP",
    "ko": "ko-KR",
    "zh-cn": "zh-CN",
    "zh-tw": "zh-TW"
}

def normalize_lang(lang_code: str) -> str:
    """Chuyển ISO-639-1 sang locale, nếu không có thì trả về nguyên bản."""
    return LANG_MAP.get(lang_code.lower(), lang_code)

def extract_text(docx_path):
    try:
        doc = Document(docx_path)

        # Lấy paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Lấy bảng (giữ nguyên dạng lưới, không mất xuống dòng trong cell)
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = "\n".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    row_data.append(cell_text)
                table_data.append(row_data)
            tables.append(table_data)

        # Lấy header/footer
        headers, footers = [], []
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                headers.extend([p.text for p in section.header.paragraphs if p.text.strip()])
            if section.footer and section.footer.paragraphs:
                footers.extend([p.text for p in section.footer.paragraphs if p.text.strip()])

        # Ghép toàn bộ text
        all_text = "\n".join(paragraphs)
        for table in tables:
            for row in table:
                all_text += "\n" + " | ".join(row)

        # Detect language (đa ngôn ngữ, trả về list)
        
        language = None
        try:
            if all_text:
                langs = detect_langs(all_text[:5000])  # lấy mẫu tối đa 5000 ký tự
                if langs:
                # lấy lang có xác suất cao nhất
                    top_lang = max(langs, key=lambda x: x.prob)
                language = normalize_lang(str(top_lang.lang))
        except Exception:
            language = None

        # Lấy metadata
        core = doc.core_properties
        metadata = {
            "file_path": os.path.abspath(docx_path),
            "title": core.title,
            "author": core.author,
            "last_modified_by": core.last_modified_by,
            "created": core.created.strftime("%Y-%m-%d %H:%M:%S") if core.created else None,
            "modified": core.modified.strftime("%Y-%m-%d %H:%M:%S") if core.modified else None,
            "pages": None,  # python-docx không hỗ trợ
            "language": language,
            "file_size": os.path.getsize(docx_path),
            "paragraphs_count": len(paragraphs),
            "tables_count": len(tables),
            "headers": headers,
            "footers": footers
        }

        return {
            "text": all_text.strip(),
            "confidence": 1.0,
            "metadata": metadata,
            "paragraphs": paragraphs,
            "tables": tables
        }

    except Exception as e:
        return {
            "text": None,
            "confidence": 0,
            "error": str(e),
            "metadata": {
                "file_path": docx_path
            }
        }

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(json.dumps({
            "text": None,
            "confidence": 0,
            "error": "Usage: python script.py <docx_path>",
            "metadata": {}
        }, ensure_ascii=False, indent=2))
        sys.exit(1)
    
    docx_path = sys.argv[1]
    result = extract_text(docx_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
